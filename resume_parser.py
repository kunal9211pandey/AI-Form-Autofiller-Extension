import re
from PyPDF2 import PdfReader
from docx import Document

class ResumeParser:
    def __init__(self):
        self.section_patterns = {
            'personal': r'(personal\s*(info|information|details)?|contact\s*(info|information|details)?)',
            'education': r'(education|academic|qualification|degree)',
            'experience': r'(experience|work\s*history|employment|professional\s*experience)',
            'skills': r'(skills|technical\s*skills|competencies|expertise)',
            'projects': r'(projects|portfolio|work\s*samples)',
            'certifications': r'(certifications?|certificates?|licenses?)',
            'summary': r'(summary|objective|profile|about\s*me)',
            'achievements': r'(achievements?|accomplishments?|awards?)',
            'languages': r'(languages?|linguistic)',
            'interests': r'(interests?|hobbies|activities)'
        }
    
    def parse(self, filepath, file_ext):
        if file_ext == 'pdf':
            text = self._parse_pdf(filepath)
        elif file_ext == 'docx':
            text = self._parse_docx(filepath)
        else:
            text = self._parse_txt(filepath)
        
        structured_data = self._extract_sections(text)
        structured_data['raw_text'] = text
        structured_data['extracted_info'] = self._extract_common_info(text)
        
        return structured_data
    
    def _parse_pdf(self, filepath):
        text = ""
        try:
            reader = PdfReader(filepath)
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        return text
    
    def _parse_docx(self, filepath):
        text = ""
        try:
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        return text
    
    def _parse_txt(self, filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    def _extract_sections(self, text):
        sections = {}
        lines = text.split('\n')
        current_section = 'general'
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            section_found = None
            
            for section_name, pattern in self.section_patterns.items():
                if re.search(pattern, line_lower) and len(line.strip()) < 50:
                    section_found = section_name
                    break
            
            if section_found:
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                current_section = section_found
                current_content = []
            else:
                if line.strip():
                    current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _extract_common_info(self, text):
        info = {}
        
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
        phones = re.findall(phone_pattern, text)
        if phones:
            info['phone'] = phones[0]
        
        linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9-]+)'
        linkedin = re.search(linkedin_pattern, text.lower())
        if linkedin:
            info['linkedin'] = linkedin.group(0)
        
        github_pattern = r'github\.com/([a-zA-Z0-9-]+)'
        github = re.search(github_pattern, text.lower())
        if github:
            info['github'] = github.group(0)
        
        lines = text.split('\n')
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if line and len(line) > 2 and len(line) < 50:
                if not re.search(r'@|http|www|\d{5,}', line):
                    if not any(keyword in line.lower() for keyword in ['resume', 'cv', 'curriculum']):
                        info['name'] = line
                        break
        
        return info
